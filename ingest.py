"""
OrgMind @ SMART — Document Ingestion v5.1
==========================================
New in v5.1:
  - Manifest tracking: already-ingested files are skipped on rerun
    (no more re-processing all 231 files every time you click Rebuild)
  - Batching: only `batch_size` NEW files are processed per call,
    so each run stays under Render's request timeout (fixes 502 errors)
  - app.py can call run_ingest() repeatedly (via st.rerun loop) until
    result["remaining"] == 0

New in v5.0:
  - Vision AI: images in PDFs and DOCX are described by Claude
  - New General CAMP subfolders added
  - Excel support retained
  - All image types: equipment photos, charts, diagrams, process flows
"""

import os
import sys
import json
import base64
import anthropic
from dotenv import load_dotenv
load_dotenv()

# Bootstrap API keys from config files
import os as _os
def _bootstrap():
    paths = ["/tmp/apikeys.txt", "/app/config.txt",
             "./config.txt", "/mount/src/orgmind-camp/config.txt"]
    for path in paths:
        if _os.path.exists(path):
            try:
                with open(path) as f:
                    for line in f:
                        line = line.strip()
                        if "=" in line and not line.startswith("#"):
                            k, v = line.split("=", 1)
                            k, v = k.strip(), v.strip()
                            if v and not v.startswith("replace_with"):
                                _os.environ.setdefault(k, v)
            except Exception:
                pass
_bootstrap()

import fitz  # PyMuPDF
from docx import Document as Docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document

try:
    import openpyxl
    EXCEL_SUPPORTED = True
except ImportError:
    EXCEL_SUPPORTED = False

CHROMA_PATH = "./chroma_db"
MANIFEST_PATH = "./ingested_manifest.json"

# ── Folder → Collection mapping ───────────────────────────────────────────────
DEPT_CONFIG = {
    "camp_legal": {
        "label": "Legal & Contracts",
        "folders": [
            "camp_documents/01_Legal_Contracts/RCA",
            "camp_documents/01_Legal_Contracts/NDA",
            "camp_documents/01_Legal_Contracts/MTA",
            "camp_documents/01_Legal_Contracts/LOA",
            "camp_documents/01_Legal_Contracts/Miscellaneous",
            "camp_documents/01_Legal_Contracts/Sub-Contracts",
            "camp_documents/01_Legal_Contracts",
        ]
    },
    "camp_staff": {
        "label": "Staff Related",
        "folders": [
            "camp_documents/02_Staff_Related/SMART-Policies",
            "camp_documents/02_Staff_Related",
        ]
    },
    "camp_research_ops": {
        "label": "Research",
        "folders": [
            "camp_documents/03_Research/Reports",
            "camp_documents/03_Research/Research-Publications-Presentations-Discussions",
            "camp_documents/03_Research",
        ]
    },
    "camp_general": {
        "label": "General CAMP",
        "folders": [
            "camp_documents/04_General_CAMP/Activity Risk Assessments",
            "camp_documents/04_General_CAMP/Project Risk Assessment",
            "camp_documents/04_General_CAMP/Safety Orientation and training",
            "camp_documents/04_General_CAMP/SOPs",
            "camp_documents/04_General_CAMP/Equipment",
            "camp_documents/04_General_CAMP/Lab_Inventory",
            "camp_documents/04_General_CAMP",
        ]
    },
}


# ── Manifest helpers ───────────────────────────────────────────────────────────
def load_manifest():
    """
    Returns dict: { filepath: {"collection": ..., "chunks": N, "mtime": ...} }
    An empty/missing manifest just means nothing has been ingested yet.
    """
    if os.path.exists(MANIFEST_PATH):
        try:
            with open(MANIFEST_PATH, "r") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}


def save_manifest(manifest):
    try:
        with open(MANIFEST_PATH, "w") as f:
            json.dump(manifest, f, indent=2)
    except Exception as e:
        print(f"  WARNING: could not save manifest: {e}")


# ── Vision AI — describe images using Claude ──────────────────────────────────
def describe_image(image_bytes, image_context="", page_num=None):
    """
    Send an image to Claude Vision and get a text description.
    Returns description string or empty string if failed.
    """
    try:
        client = anthropic.Anthropic(
            api_key=os.environ.get("ANTHROPIC_API_KEY", ""))

        # Encode image
        img_b64 = base64.standard_b64encode(image_bytes).decode("utf-8")

        # Detect image type
        if image_bytes[:4] == b'\x89PNG':
            media_type = "image/png"
        elif image_bytes[:2] == b'\xff\xd8':
            media_type = "image/jpeg"
        elif image_bytes[:4] == b'GIF8':
            media_type = "image/gif"
        else:
            media_type = "image/png"  # default

        # Skip tiny images (likely icons/logos) — under 5KB
        if len(image_bytes) < 5000:
            return ""

        context_hint = ""
        if image_context:
            context_hint = f"This image appears in: {image_context}. "
        if page_num:
            context_hint += f"It is on page {page_num}. "

        prompt = (
            f"{context_hint}"
            "Describe this image in detail for a research institute knowledge base. "
            "Focus on: equipment names and model numbers, locations mentioned, "
            "data shown in charts or graphs (include specific values if visible), "
            "process steps in diagrams, safety information, key measurements, "
            "and any text visible in the image. "
            "Be specific and factual. If this is a chart, describe the data. "
            "If this is equipment, describe what it is and any identifying details. "
            "If this is a diagram, describe what process or system it shows. "
            "Keep the description under 200 words."
        )

        response = client.messages.create(
            model="claude-opus-4-5-20251101",
            max_tokens=300,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": img_b64,
                        },
                    },
                    {"type": "text", "text": prompt}
                ],
            }]
        )
        return response.content[0].text.strip()

    except Exception:
        return ""


# ── Document readers ──────────────────────────────────────────────────────────
def read_pdf(path, fname, enable_vision=True):
    """
    Extract text AND describe images from a PDF.
    Returns combined text with image descriptions inline.
    """
    doc = fitz.open(path)
    all_content = []
    image_count = 0

    for page_num, page in enumerate(doc, 1):
        page_text = page.get_text()
        if page_text.strip():
            all_content.append(f"[Page {page_num}]\n{page_text}")

        if enable_vision:
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]

                    description = describe_image(
                        image_bytes,
                        image_context=f"{fname}, page {page_num}",
                        page_num=page_num
                    )

                    if description:
                        image_count += 1
                        all_content.append(
                            f"\n[IMAGE on page {page_num}]: {description}\n"
                        )
                except Exception:
                    pass

    doc.close()

    if image_count > 0:
        print(f"      → {image_count} image(s) described by Vision AI")

    return "\n".join(all_content)


def read_docx(path, fname, enable_vision=True):
    """
    Extract text AND describe images from a DOCX file.
    """
    doc = Docx(path)
    all_content = []
    image_count = 0

    for para in doc.paragraphs:
        if para.text.strip():
            all_content.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = "  |  ".join(
                cell.text.strip() for cell in row.cells
                if cell.text.strip()
            )
            if row_text:
                all_content.append(row_text)

    if enable_vision:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_bytes = rel.target_part.blob
                    description = describe_image(
                        image_bytes,
                        image_context=fname
                    )
                    if description:
                        image_count += 1
                        all_content.append(
                            f"\n[IMAGE in document]: {description}\n"
                        )
                except Exception:
                    pass

    if image_count > 0:
        print(f"      → {image_count} image(s) described by Vision AI")

    return "\n".join(all_content)


def read_excel(path):
    """Convert Excel to searchable text."""
    if not EXCEL_SUPPORTED:
        return ""
    try:
        wb = openpyxl.load_workbook(path, data_only=True)
        lines = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"[Sheet: {sheet_name}]")
            for row in ws.iter_rows():
                row_values = [
                    str(cell.value).strip()
                    for cell in row
                    if cell.value is not None
                ]
                if row_values:
                    lines.append("  |  ".join(row_values))
        return "\n".join(lines)
    except Exception as e:
        raise Exception(f"Could not read Excel: {e}")


# ── File discovery (manifest-aware) ───────────────────────────────────────────
def discover_pending_files(manifest):
    """
    Walks all configured folders and returns a list of
    (filepath, fname, ext, collection_name) tuples for files
    NOT already in the manifest. Order is stable (sorted) so
    batches are deterministic across repeated calls.
    """
    pending = []
    seen_paths = set()

    for collection_name, config in DEPT_CONFIG.items():
        for folder in config["folders"]:
            if not os.path.exists(folder):
                continue
            for fname in sorted(os.listdir(folder)):
                ext = fname.lower().split('.')[-1]
                if ext not in ('pdf', 'docx', 'xlsx', 'xls'):
                    continue
                if fname.startswith(("~", ".")):
                    continue

                fpath = os.path.join(folder, fname)
                if fpath in seen_paths:
                    continue  # same file reachable via two folder entries (e.g. parent + subfolder listing)
                seen_paths.add(fpath)

                if fpath in manifest:
                    continue  # already ingested in a previous run/batch

                pending.append((fpath, fname, ext, collection_name))

    return pending


def process_file(fpath, fname, ext, enable_vision):
    """
    Reads a single file and returns a langchain Document, or None if
    skipped (too short / scanned / unreadable).
    Raises on hard errors so the caller can log and continue.
    """
    if ext == 'pdf':
        text = read_pdf(fpath, fname, enable_vision)
    elif ext in ('xlsx', 'xls'):
        text = read_excel(fpath)
    else:
        text = read_docx(fpath, fname, enable_vision)

    if len(text.strip()) < 100:
        return None

    fname_upper = fname.upper()
    doc_type = "GENERAL"
    for t in ["RCA", "NDA", "MTA", "LOA"]:
        if t in fname_upper:
            doc_type = t
            break

    return Document(
        page_content=text,
        metadata={
            "source": fname,
            "doc_type": doc_type,
            "filepath": fpath
        }
    )


# ── Main ingestion (batched, manifest-aware) ──────────────────────────────────
def run_ingest(enable_vision=True, batch_size=10):
    """
    Process up to `batch_size` NEW (not-yet-ingested) files, embed them,
    and update the manifest. Safe to call repeatedly — each call only
    touches files that haven't been ingested yet, so a 502/timeout never
    loses progress on files that already completed.

    Returns:
      {
        "success": bool,
        "processed": int,       # files processed in THIS call
        "remaining": int,       # files still pending after this call
        "total_chunks": int,    # chunks added in THIS call
        "details": {...},       # per-collection breakdown for this call
        "vision_enabled": bool,
        "error": str (only if success=False)
      }
    """
    if enable_vision:
        key = os.environ.get("ANTHROPIC_API_KEY", "")
        if not key or key.startswith("paste") or key.startswith("replace"):
            print("  WARNING: Anthropic API key not found — "
                  "running without Vision AI")
            enable_vision = False

    manifest = load_manifest()

    try:
        embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
    except Exception as e:
        return {"success": False, "error": str(e),
                "processed": 0, "remaining": 0,
                "total_chunks": 0, "details": {}}

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1500,
        chunk_overlap=400,
        separators=["\n\n", "\n", ". ", " "]
    )

    pending = discover_pending_files(manifest)
    total_pending = len(pending)
    this_batch = pending[:batch_size]

    print(f"\n[Batch] {len(this_batch)} file(s) this run, "
          f"{total_pending} total pending before this run")

    # Group this batch's docs by collection so we can write to the right
    # Chroma collection, while still respecting the overall batch_size cap.
    docs_by_collection = {name: [] for name in DEPT_CONFIG}
    processed = 0
    details = {}

    for fpath, fname, ext, collection_name in this_batch:
        label = DEPT_CONFIG[collection_name]["label"]
        try:
            doc = process_file(fpath, fname, ext, enable_vision)
            if doc is None:
                print(f"    SKIPPED (too short/scanned): {fname}")
                # Mark as seen so we don't retry it every batch forever
                manifest[fpath] = {
                    "collection": collection_name,
                    "chunks": 0,
                    "status": "skipped"
                }
                continue

            docs_by_collection[collection_name].append(doc)
            processed += 1
            print(f"    Loaded ({label}): {fname}")

        except Exception as e:
            print(f"    ERROR {fname}: {e}")
            # Don't add to manifest — will be retried next batch

    # Embed + write each collection's new docs
    for collection_name, new_docs in docs_by_collection.items():
        if not new_docs:
            continue
        label = DEPT_CONFIG[collection_name]["label"]
        chunks = splitter.split_documents(new_docs)
        try:
            Chroma.from_documents(
                documents=chunks,
                embedding=embeddings,
                collection_name=collection_name,
                persist_directory=CHROMA_PATH
            )
            details[label] = {"docs": len(new_docs), "chunks": len(chunks)}

            # Only mark files as ingested AFTER a successful Chroma write
            for doc in new_docs:
                manifest[doc.metadata["filepath"]] = {
                    "collection": collection_name,
                    "chunks": len(chunks),  # approximate (per-collection batch)
                    "status": "ingested"
                }

        except Exception as e:
            details[label] = {"docs": 0, "chunks": 0, "error": str(e)}
            print(f"    ERROR writing collection {label} to Chroma: {e}")

    save_manifest(manifest)

    remaining = max(0, total_pending - processed)
    total_chunks = sum(d.get("chunks", 0) for d in details.values())

    return {
        "success": True,
        "processed": processed,
        "remaining": remaining,
        "total_chunks": total_chunks,
        "details": details,
        "vision_enabled": enable_vision
    }


def reset_manifest():
    """
    Wipes the manifest only (does NOT touch Chroma DB).
    Use this if you intentionally want to force a full re-ingest,
    e.g. after changing chunking strategy or fixing a bug in extraction.
    """
    if os.path.exists(MANIFEST_PATH):
        os.remove(MANIFEST_PATH)


if __name__ == "__main__":
    print("\n" + "="*60)
    print("  OrgMind @ CAMP — Document Ingestion v5.1 (batched)")
    print("="*60)
    result = run_ingest()
    if not result["success"]:
        print(f"\n ERROR: {result['error']}")
    else:
        vision = "✅ ON" if result.get("vision_enabled") else "⚠️ OFF"
        print(f"\n  Batch done: {result['processed']} file(s) processed")
        print(f"              {result['remaining']} file(s) still pending")
        print(f"              {result['total_chunks']} chunks added this batch")
        print(f"              Vision AI: {vision}")
    print("="*60 + "\n")
